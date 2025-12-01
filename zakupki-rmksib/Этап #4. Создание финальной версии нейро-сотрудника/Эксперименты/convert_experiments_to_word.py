"""
Скрипт для конвертации Эксперименты_нейро-сотрудник.md в Word документ

Использование:
    python convert_experiments_to_word.py
"""

import os
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Ошибка: Не установлена библиотека python-docx")
    print("Установите её командой: pip install python-docx")
    sys.exit(1)


def parse_table(line: str) -> list:
    """Парсинг строки таблицы Markdown"""
    # Убираем начальные и конечные |
    line = line.strip().strip('|')
    # Разбиваем по |
    cells = [cell.strip() for cell in line.split('|')]
    return cells


def markdown_to_word(md_file: str, output_file: str):
    """Конвертация Markdown файла в Word документ с поддержкой таблиц"""
    
    # Читаем Markdown файл
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Создаем новый Word документ
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Обработка содержимого
    lines = content.split('\n')
    i = 0
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Пропускаем пустые строки (но не внутри таблицы)
        if not line and not in_table:
            i += 1
            continue
        
        # Обработка таблиц
        if '|' in line and not line.startswith('```'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Пропускаем разделитель таблицы (---)
            if re.match(r'^\|[\s\-:]+\|', line):
                i += 1
                continue
            
            cells = parse_table(line)
            if cells:
                table_data.append(cells)
            i += 1
            
            # Проверяем, закончилась ли таблица
            if i < len(lines):
                next_line = lines[i].strip()
                if not next_line or (not '|' in next_line and not next_line.startswith('```')):
                    # Создаем таблицу
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Light Grid Accent 1'
                        
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_data in enumerate(row_data):
                                if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                                    cell = table.rows[row_idx].cells[col_idx]
                                    cell.text = cell_data
                                    # Заголовок таблицы жирным
                                    if row_idx == 0:
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.bold = True
                        
                        doc.add_paragraph()  # Пустая строка после таблицы
                        table_data = []
                        in_table = False
            continue
        
        # Если мы были в таблице, но строка не содержит |, выходим
        if in_table:
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                doc.add_paragraph()
                table_data = []
            in_table = False
        
        # Заголовки
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            i += 1
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            i += 1
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
            i += 1
        elif line.startswith('```'):
            # Блок кода
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            
            # Добавляем код как параграф с моноширинным шрифтом
            if code_lines:
                code_para = doc.add_paragraph('\n'.join(code_lines))
                code_para.style = 'No Spacing'
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
        elif line.startswith('- ') or line.startswith('* '):
            # Маркированный список
            doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        elif re.match(r'^\d+\.\s', line):
            # Нумерованный список
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
            i += 1
        else:
            # Обычный текст
            para = doc.add_paragraph()
            
            # Обрабатываем форматирование: **жирный**, `код`, [OK], [ERROR]
            text = line
            
            # Заменяем эмодзи-маркеры на текст
            text = text.replace('✅', '[OK]')
            text = text.replace('❌', '[ERROR]')
            text = text.replace('⚠️', '[WARNING]')
            text = text.replace('💡', '[TIP]')
            text = text.replace('📊', '[CHART]')
            text = text.replace('🎯', '[TARGET]')
            text = text.replace('📋', '[LIST]')
            text = text.replace('🚀', '[ROCKET]')
            
            # Обрабатываем жирный текст **text**
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    para.add_run(part)
            
            i += 1
    
    # Сохраняем документ
    doc.save(output_file)
    print(f"Документ успешно создан: {output_file}")


def main():
    # Определяем пути
    script_dir = Path(__file__).parent
    md_file = script_dir / "Эксперименты_нейро-сотрудник.md"
    output_file = script_dir / "Эксперименты_нейро-сотрудник.docx"
    
    if not md_file.exists():
        print(f"Ошибка: Файл {md_file} не найден")
        sys.exit(1)
    
    print(f"Конвертация {md_file} в {output_file}...")
    markdown_to_word(str(md_file), str(output_file))
    print("Конвертация завершена!")


if __name__ == "__main__":
    main()

