"""
Скрипт для конвертации NEURO_EMPLOYEE_REPORT.md в Word документ

Использование:
    python convert_to_word.py
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Ошибка: Не установлена библиотека python-docx")
    print("Установите её командой: pip install python-docx")
    sys.exit(1)


def markdown_to_word(md_file: str, output_file: str):
    """Конвертация Markdown файла в Word документ"""
    
    # Читаем Markdown файл
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Создаем новый Word документ
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Обработка содержимого
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Пропускаем пустые строки
        if not line:
            i += 1
            continue
        
        # Заголовки
        if line.startswith('# '):
            # H1
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
        elif line.startswith('## '):
            # H2
            doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            # H3
            doc.add_heading(line[4:], level=3)
            i += 1
        elif line.startswith('#### '):
            # H4
            doc.add_heading(line[5:], level=4)
            i += 1
        elif line.startswith('---'):
            # Горизонтальная линия
            doc.add_paragraph('─' * 50)
            i += 1
        elif line.startswith('```'):
            # Блок кода
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Пропускаем закрывающий ```
            
            # Добавляем код как параграф с моноширинным шрифтом
            code_para = doc.add_paragraph('\n'.join(code_lines))
            code_para.style = 'No Spacing'
            for run in code_para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        elif line.startswith('- ') or line.startswith('* '):
            # Маркированный список
            doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            # Нумерованный список
            doc.add_paragraph(line[3:], style='List Number')
            i += 1
        else:
            # Обычный текст
            # Обрабатываем жирный текст
            para = doc.add_paragraph()
            parts = line.split('**')
            for j, part in enumerate(parts):
                run = para.add_run(part)
                if j % 2 == 1:  # Нечетные индексы - жирный текст
                    run.bold = True
            
            i += 1
    
    # Сохраняем документ
    doc.save(output_file)
    print(f"Документ успешно создан: {output_file}")


def main():
    # Определяем пути
    script_dir = Path(__file__).parent
    md_file = script_dir / "NEURO_EMPLOYEE_REPORT.md"
    output_file = script_dir / "NEURO_EMPLOYEE_REPORT.docx"
    
    if not md_file.exists():
        print(f"Ошибка: Файл {md_file} не найден")
        sys.exit(1)
    
    print(f"Конвертация {md_file} в {output_file}...")
    markdown_to_word(str(md_file), str(output_file))
    print("Конвертация завершена!")


if __name__ == "__main__":
    main()

